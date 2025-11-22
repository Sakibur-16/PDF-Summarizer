# document_processor.py
import re
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
from docx import Document
from pathlib import Path
from config import Config
from logger_module import DocumentLogger

class DocumentProcessor:
    """Process PDF and DOCX files with OCR support"""
    
    def __init__(self, file_path, language='en', logger=None):
        self.file_path = Path(file_path)
        self.language = language
        self.logger = logger
        self.text = ""
        self.chapters = []
        self.metadata = {}
        
        # Set Tesseract path
        pytesseract.pytesseract.tesseract_cmd = Config.TESSERACT_CMD
        
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
    
    def process(self):
        """Main processing method"""
        file_extension = self.file_path.suffix.lower()
        
        if self.logger:
            file_size = self.file_path.stat().st_size
            self.logger.log_processing_start(str(self.file_path), file_size, self.language)
        
        if file_extension == '.pdf':
            self.text = self._process_pdf()
        elif file_extension in ['.docx', '.doc']:
            self.text = self._process_docx()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Extract chapters
        self.chapters = self._extract_chapters()
        
        # Extract metadata
        self.metadata = self._extract_metadata()
        
        if self.logger:
            self.logger.log_event(
                f"Document processed: {len(self.text)} characters, {len(self.chapters)} chapters found",
                "INFO"
            )
        
        return {
            'text': self.text,
            'chapters': self.chapters,
            'metadata': self.metadata
        }
    
    def _process_pdf(self):
        """Process PDF with OCR fallback"""
        try:
            # Try extracting text directly
            reader = PdfReader(str(self.file_path))
            text = ""
            
            for page in reader.pages:
                page_text = page.extract_text()
                text += page_text + "\n\n"
            
            # If very little text extracted, use OCR
            if len(text.strip()) < 100:
                if self.logger:
                    self.logger.log_event("Low text extraction, switching to OCR", "INFO")
                text = self._process_pdf_with_ocr()
            else:
                if self.logger:
                    self.logger.log_event(f"Text extracted from {len(reader.pages)} pages", "INFO")
            
            return text
        
        except Exception as e:
            if self.logger:
                self.logger.log_event(f"PDF text extraction failed: {str(e)}", "WARNING")
            
            # Fallback to OCR
            return self._process_pdf_with_ocr()
    
    def _process_pdf_with_ocr(self):
        """Process PDF using OCR"""
        try:
            if self.logger:
                self.logger.log_event("Starting OCR processing", "INFO")
            
            # Convert PDF to images
            images = convert_from_path(str(self.file_path))
            
            text = ""
            tesseract_lang = Config.TESSERACT_LANG_MAP.get(self.language, 'eng')
            
            for i, image in enumerate(images):
                if self.logger and i % 5 == 0:
                    self.logger.log_event(f"Processing page {i+1}/{len(images)} with OCR", "DEBUG")
                
                # Perform OCR
                page_text = pytesseract.image_to_string(image, lang=tesseract_lang)
                text += page_text + "\n\n"
            
            if self.logger:
                self.logger.log_ocr_operation(len(images), success=True)
            
            return text
        
        except Exception as e:
            if self.logger:
                self.logger.log_event(f"OCR failed: {str(e)}", "ERROR")
                self.logger.log_ocr_operation(0, success=False)
            raise Exception(f"OCR processing failed: {str(e)}")
    
    def _process_docx(self):
        """Process DOCX files"""
        try:
            doc = Document(str(self.file_path))
            text = ""
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if self.logger:
                self.logger.log_event(f"DOCX processed: {len(doc.paragraphs)} paragraphs", "INFO")
            
            return text
        
        except Exception as e:
            if self.logger:
                self.logger.log_event(f"DOCX processing failed: {str(e)}", "ERROR")
            raise Exception(f"DOCX processing failed: {str(e)}")
    
    def _extract_chapters(self):
        """Extract chapters from text"""
        chapters = []
        
        # Common chapter patterns
        chapter_patterns = [
            r'(?:Chapter|CHAPTER|অধ্যায়|الفصل)\s*(\d+|[IVX]+)\s*[:\-–—]?\s*(.+?)(?=\n)',
            r'(?:Part|PART|অংশ|الجزء)\s*(\d+|[IVX]+)\s*[:\-–—]?\s*(.+?)(?=\n)',
            r'^(\d+)\.\s+(.+?)(?=\n)',
            r'^([IVX]+)\.\s+(.+?)(?=\n)',
        ]
        
        # Try each pattern
        for pattern in chapter_patterns:
            matches = re.finditer(pattern, self.text, re.MULTILINE | re.IGNORECASE)
            chapter_positions = []
            
            for match in matches:
                chapter_num = match.group(1)
                chapter_title = match.group(2).strip()
                start_pos = match.start()
                
                chapter_positions.append({
                    'number': chapter_num,
                    'title': chapter_title,
                    'start': start_pos
                })
            
            if chapter_positions:
                # Extract chapter content
                for i, chapter in enumerate(chapter_positions):
                    end_pos = chapter_positions[i+1]['start'] if i+1 < len(chapter_positions) else len(self.text)
                    content = self.text[chapter['start']:end_pos].strip()
                    
                    chapters.append({
                        'number': chapter['number'],
                        'title': chapter['title'],
                        'content': content,
                        'word_count': len(content.split())
                    })
                    
                    if self.logger:
                        self.logger.log_chapter_processed(
                            f"{chapter['number']}: {chapter['title']}", 
                            len(content.split())
                        )
                
                break
        
        # If no chapters found, treat whole document as one chapter
        if not chapters:
            chapters.append({
                'number': '1',
                'title': 'Full Document',
                'content': self.text,
                'word_count': len(self.text.split())
            })
            
            if self.logger:
                self.logger.log_event("No chapters detected, treating as single document", "INFO")
        
        return chapters
    
    def _extract_metadata(self):
        """Extract metadata from document"""
        metadata = {
            'file_name': self.file_path.name,
            'file_size': self.file_path.stat().st_size,
            'total_words': len(self.text.split()),
            'total_characters': len(self.text),
            'chapter_count': len(self.chapters)
        }
        
        # Try to find title (first significant line)
        lines = self.text.split('\n')
        for line in lines:
            if len(line.strip()) > 10 and len(line.strip()) < 200:
                metadata['title'] = line.strip()
                break
        
        return metadata